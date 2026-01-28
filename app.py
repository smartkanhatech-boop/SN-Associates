import streamlit as st
import pandas as pd
from datetime import datetime, date
from fpdf import FPDF
import json
import os
import io
import base64
import math
import random
import re

# Word Document Library
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Optional: Number to Words
try:
    from num2words import num2words
    HAS_NUM2WORDS = True
except ImportError:
    HAS_NUM2WORDS = False

# --- PAGE CONFIG ---
st.set_page_config(page_title="SN Associates Billing", layout="wide", page_icon="🏗️")

# --- AUTHENTICATION ---
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False

def check_login():
    if st.session_state.username == 'chaitanyababu2603' and st.session_state.password == 'myson@2501':
        st.session_state.authenticated = True
        st.session_state.login_error = False
    else:
        st.session_state.login_error = True

if not st.session_state.authenticated:
    st.markdown("<h2 style='text-align: center; color: #000080;'>🔒 SN Associates Login</h2>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1,2,1])
    with c2:
        st.text_input("Username", key="username")
        st.text_input("Password", type="password", key="password")
        st.button("Login", on_click=check_login, use_container_width=True)
        if st.session_state.get('login_error'):
            st.error("Access Denied")
    st.stop()

# --- CONSTANTS ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_FULL_PATH = os.path.join(BASE_DIR, "Logo.png")
DB_FILE = os.path.join(BASE_DIR, "sn_billing_db.json")

# --- COMPANY DETAILS ---
COMPANY_NAME = "SN ASSOCIATES"
COMPANY_PHONE = "+91 91318 71696"
COMPANY_GSTIN = "23AMAPT7272P1ZG"
COMPANY_ADDRESS = "Chhatarpur, MP"
SIGNATORY_NAME = "(Dheerendra Tiwari)"

# --- BANK DETAILS ---
BANK_DETAILS = {
    "bank": "HDFC BANK",
    "ac": "99999131871696",
    "ifsc": "HDFC0004849",
    "name": "SN Associates"
}

# --- WORK CATALOG (UPDATED: Categories moved to Description Items) ---
WORK_CATALOG = [
    "Site Visit",
    "Architecture & Design",
    "Structural Design",
    "Electrical & Plumbing",
    "2D & 3D",
    "Elevation Drawing",
    "Landscape Drawing",
    "Vastu consultancy",
    "Supply",
    "Walkthrough"
]

GST_RATES = {"0%": 0.00, "5%": 0.05, "12%": 0.12, "18%": 0.18}

# --- SESSION STATE ---
if 'invoice_data' not in st.session_state:
    st.session_state.invoice_data = {
        "items": [],
        "schedule": [],
        "meta": {"terms": """- 30% Advance prior to initiation of the work.
- 2 extra changes will be provided free of cost. Further changes will be chargeable as per requirement.
- Site visit will be chargeable (unless specified above).
- The project is to be completed within six months. In case of delay, the agreed price will be revised by 10% for every additional two months."""}
    }

if 'builder_c_name' not in st.session_state: st.session_state.builder_c_name = ""
if 'builder_c_mob' not in st.session_state: st.session_state.builder_c_mob = ""
if 'builder_c_addr' not in st.session_state: st.session_state.builder_c_addr = ""
if 'builder_dtype_idx' not in st.session_state: st.session_state.builder_dtype_idx = 0

if 'schedule_df' not in st.session_state:
    st.session_state.schedule_df = pd.DataFrame(columns=["Stage", "Amount", "Date"])

# --- DB & HELPERS ---
def save_db():
    with open(DB_FILE, 'w') as f: json.dump(st.session_state.db, f)

def load_db():
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, 'r') as f: db = json.load(f)
        except: db = {"invoices": [], "quotations": [], "payments": []}
    else:
        db = {"invoices": [], "quotations": [], "payments": []}
    
    needs_save = False
    for rec in db.get('invoices', []):
        if 'id' not in rec:
            rec['id'] = f"LEGACY_{random.randint(1000,9999)}_{int(datetime.now().timestamp())}"
            needs_save = True
        if 'status' not in rec:
            rec['status'] = "Pending"
            needs_save = True
        if 'invoice_no' not in rec:
            rec['invoice_no'] = rec['id'] 
            needs_save = True
            
    if 'payments' not in db:
        db['payments'] = []
        needs_save = True
        
    if needs_save:
        with open(DB_FILE, 'w') as f: json.dump(db, f)
    return db

if 'db' not in st.session_state:
    st.session_state.db = load_db()

def sanitize_text(text):
    if not isinstance(text, str): text = str(text)
    replacements = {'₹': 'Rs. ', '\u20b9': 'Rs. ', '•': '-', '–': '-', '—': '-', '“': '"', '”': '"', '‘': "'", '’': "'", '…': '...'}
    for old, new in replacements.items(): text = text.replace(old, new)
    return text.encode('latin-1', 'ignore').decode('latin-1')

def number_to_words_safe(amount):
    if HAS_NUM2WORDS:
        try:
            txt = num2words(amount, lang='en_IN').title()
            return sanitize_text(txt) + " Only"
        except: return "Check Amount"
    return f"{amount} (in words)"

def calculate_totals(items, gst_rate_key):
    sub = sum(item['qty'] * item['rate'] for item in items)
    rate = GST_RATES.get(gst_rate_key, 0.0)
    gst = sub * rate
    grand = sub + gst
    return sub, gst, grand

# --- ID GENERATOR ---
def generate_next_id(doc_type, date_obj):
    year = str(date_obj.year)
    prefix = "INV" if doc_type == "FINAL BILL" else "QUOT"
    target_format = f"{prefix}-{year}-"
    
    record_list = st.session_state.db['invoices'] if doc_type == "FINAL BILL" else st.session_state.db['quotations']
    max_seq = 0
    for rec in record_list:
        rec_no = rec.get('invoice_no') if doc_type == "FINAL BILL" else rec.get('quotation_no')
        if rec_no and rec_no.startswith(target_format):
            try:
                seq_part = rec_no.split('-')[-1]
                seq = int(seq_part)
                if seq > max_seq: max_seq = seq
            except: pass
    new_seq = max_seq + 1
    return f"{target_format}{new_seq:03d}"

# --- RECEIPT PDF ---
class ReceiptPDF(FPDF):
    def header(self): pass
    def footer(self): pass

def generate_receipt_bytes(payment_data):
    try:
        pdf = ReceiptPDF(format='A5', orientation='L')
        pdf.add_page()
        pdf.set_draw_color(0,0,0); pdf.rect(5, 5, 200, 138)
        if os.path.exists(LOGO_FULL_PATH): 
            try: pdf.image(LOGO_FULL_PATH, 10, 10, 25)
            except: pass
        pdf.set_y(10); pdf.set_font('Times', 'B', 16)
        pdf.set_text_color(0, 0, 128) 
        pdf.cell(0, 8, sanitize_text(COMPANY_NAME), 0, 1, 'R')
        pdf.set_text_color(0,0,0); pdf.set_font('Times', '', 9)
        pdf.cell(0, 5, sanitize_text(COMPANY_ADDRESS), 0, 1, 'R')
        pdf.cell(0, 5, sanitize_text(f"Ph: {COMPANY_PHONE}"), 0, 1, 'R')
        pdf.ln(10); pdf.set_font('Times', 'B', 14); pdf.cell(0, 10, "PAYMENT RECEIPT", 0, 1, 'C'); pdf.ln(5)
        pdf.set_font('Times', '', 12); pdf.set_x(20)
        pdf.write(8, "Received with thanks from  ")
        pdf.set_font('Times', 'B', 14); pdf.write(8, sanitize_text(payment_data['client_name']))
        pdf.set_font('Times', '', 12); pdf.write(8, "\n\n")
        pdf.set_x(20)
        text = (f"The sum of  Rs. {payment_data['amount']:,.2f}\n"
                f"({number_to_words_safe(payment_data['amount'])})\n\n"
                f"Payment Date:  {payment_data['date']}\n"
                f"Payment Mode:  {payment_data['mode']}\n"
                f"Ref Invoice Date:  {payment_data.get('invoice_date', 'N/A')}")
        pdf.multi_cell(0, 8, sanitize_text(text))
        pdf.set_y(-35); pdf.set_x(130) 
        pdf.set_font('Times', 'B', 10)
        pdf.cell(60, 5, sanitize_text(SIGNATORY_NAME), 0, 1, 'C')
        pdf.set_x(130)
        pdf.cell(60, 5, "AUTHORIZED SIGNATORY", 0, 0, 'C')
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e: return None

# --- BILL PDF ---
class PDF(FPDF):
    def header(self): pass
    def footer(self): pass

def calculate_page_height(data, schedule_list):
    min_height = 297; required_height = 160 
    for item in data['items']:
        desc_len = len(item['desc']); lines = math.ceil(desc_len / 45); row_h = max(8, lines * 5)
        required_height += row_h
    if schedule_list: required_height += 20 + (len(schedule_list) * 8)
    term_lines = data['meta']['terms'].count('\n') + 3
    required_height += (term_lines * 5)
    return max(min_height, required_height)

def generate_pdf_bytes(data, gst_rate_key, hide_gst, schedule_list, doc_no):
    try:
        page_h = calculate_page_height(data, schedule_list)
        pdf = PDF(unit='mm', format=(210, page_h))
        pdf.set_auto_page_break(False); pdf.set_margins(10, 10, 10); pdf.add_page()
        
        if os.path.exists(LOGO_FULL_PATH): 
            try: pdf.image(LOGO_FULL_PATH, 10, 10, 35)
            except: pass
        
        pdf.set_xy(110, 12); pdf.set_font('Times', 'B', 22)
        pdf.set_text_color(0, 0, 128)
        pdf.cell(90, 8, sanitize_text(COMPANY_NAME), 0, 1, 'R')
        pdf.set_xy(110, 20); pdf.set_text_color(0, 0, 0); pdf.set_font('Times', '', 10)
        pdf.cell(90, 5, sanitize_text(COMPANY_ADDRESS), 0, 1, 'R')
        pdf.set_xy(110, 25); pdf.cell(90, 5, sanitize_text(f"Ph: {COMPANY_PHONE}"), 0, 1, 'R')
        if not hide_gst: pdf.set_xy(110, 30); pdf.cell(90, 5, sanitize_text(f"GST: {COMPANY_GSTIN}"), 0, 1, 'R')
            
        pdf.set_draw_color(0, 0, 0); pdf.line(10, 50, 200, 50)
        
        pdf.set_y(55); pdf.set_font('Times', 'B', 16)
        display_type = "BILL" if data['meta']['type'] == "FINAL BILL" else data['meta']['type']
        pdf.cell(0, 8, sanitize_text(display_type), 0, 1, 'C')
        
        y_info = 68
        pdf.set_xy(10, y_info); pdf.set_font('Times', 'B', 10); pdf.cell(90, 5, "DOCUMENT DETAILS:", 0, 1)
        pdf.set_font('Times', '', 10); pdf.set_x(10)
        
        lbl = "Invoice No" if data['meta']['type'] == "FINAL BILL" else "Quotation No"
        pdf.cell(90, 5, sanitize_text(f"{lbl}: {doc_no}"), 0, 1)
        
        pdf.set_x(10); pdf.cell(90, 5, sanitize_text(f"Date: {data['meta']['date']}"), 0, 1)
        
        pdf.set_xy(110, y_info); pdf.set_font('Times', 'B', 10); pdf.cell(90, 5, "TO CLIENT:", 0, 1)
        pdf.set_xy(110, y_info + 6); pdf.set_font('Times', 'B', 12)
        pdf.cell(90, 6, sanitize_text(data['client']['name']), 0, 1)
        pdf.set_xy(110, y_info + 12); pdf.set_font('Times', '', 10)
        details = f"{data['client']['phone']}\n{data['client']['address']}"
        pdf.multi_cell(90, 5, sanitize_text(details))
        
        y_table_start = max(pdf.get_y(), y_info + 25) + 5
        pdf.set_xy(10, y_table_start)
        
        # CHANGED: 4 Columns (Removed Category), Increased Description Width
        cols = [120, 20, 25, 25] 
        pdf.set_fill_color(240, 240, 240); pdf.set_font('Times', 'B', 10)
        headers = ["Description", "Qty", "Rate (Rs.)", "Amount (Rs.)"]
        for i, h in enumerate(headers):
            align = 'L' if i == 0 else 'R'
            pdf.cell(cols[i], 8, h, 1, 0, align, 1)
        pdf.ln()
        
        pdf.set_font('Times', '', 10)
        sub, gst, grand = calculate_totals(data['items'], gst_rate_key)
        if hide_gst: gst=0; grand=sub

        for item in data['items']:
            x = pdf.get_x(); y = pdf.get_y()
            desc_txt = sanitize_text(item['desc'])
            lines = math.ceil(len(desc_txt) / 65) + desc_txt.count('\n') # Adjusted char limit for wider col
            row_h = max(8, lines * 5)
            
            # Col 1: Description (Merged Category items go here)
            pdf.set_xy(x, y); pdf.cell(cols[0], row_h, "", 1, 0, 'L')
            
            # Col 2: Qty
            pdf.set_xy(x + cols[0], y); pdf.cell(cols[1], row_h, sanitize_text(f"{item['qty']} {item['unit']}"), 1, 0, 'R')
            
            # Col 3: Rate
            pdf.set_xy(x + cols[0] + cols[1], y); pdf.cell(cols[2], row_h, f"{item['rate']:.2f}", 1, 0, 'R')
            
            # Col 4: Amount
            pdf.set_xy(x + cols[0] + cols[1] + cols[2], y); pdf.cell(cols[3], row_h, f"{item['qty']*item['rate']:.2f}", 1, 0, 'R')
            
            # Print Description Text
            pdf.set_xy(x, y)
            pdf.multi_cell(cols[0], 5, desc_txt, 0, 'L')
            
            pdf.set_y(y + row_h)

        pdf.ln(2)
        def print_total(label, val, bold=False):
            if bold: pdf.set_font('Times', 'B', 11)
            else: pdf.set_font('Times', '', 10)
            pdf.cell(165, 6, label, 0, 0, 'R'); pdf.cell(25, 6, f"Rs. {val:,.2f}", 0, 1, 'R')

        print_total("Subtotal:", sub)
        if not hide_gst: print_total(f"GST ({gst_rate_key}):", gst)
        print_total("Grand Total:", grand, bold=True)
        pdf.ln(2); pdf.set_font('Times', 'I', 10)
        pdf.cell(0, 6, number_to_words_safe(grand), 0, 1, 'R')
        
        if schedule_list and len(schedule_list) > 0:
            pdf.ln(8); pdf.set_font('Times', 'B', 10); pdf.cell(0, 6, "PAYMENT SCHEDULE:", 0, 1, 'L')
            pdf.set_fill_color(245, 245, 245)
            pdf.cell(80, 6, "Stage", 1, 0, 'L', 1); pdf.cell(40, 6, "Amount", 1, 0, 'C', 1); pdf.cell(70, 6, "Date", 1, 1, 'L', 1)
            pdf.set_font('Times', '', 9)
            for r in schedule_list:
                pdf.cell(80, 6, sanitize_text(str(r.get("Stage",""))), 1)
                pdf.cell(40, 6, sanitize_text(str(r.get("Amount",""))), 1, 0, 'C')
                pdf.cell(70, 6, sanitize_text(str(r.get("Date",""))), 1, 1)

        pdf.ln(8); pdf.set_font('Times', 'B', 10); pdf.cell(0, 6, "TERMS & CONDITIONS:", 0, 1, 'L')
        pdf.set_font('Times', '', 10); pdf.multi_cell(0, 5, sanitize_text(data['meta']['terms']))
        
        pdf.ln(10); pdf.set_draw_color(0, 0, 0); pdf.line(10, pdf.get_y(), 200, pdf.get_y()); pdf.ln(5)
        y_foot = pdf.get_y(); pdf.set_font('Times', 'B', 10); pdf.cell(90, 5, "ACCOUNT DETAILS", 0, 1, 'L')
        pdf.set_font('Times', '', 9)
        pdf.cell(90, 5, sanitize_text(f"BANK: {BANK_DETAILS['bank']}"), 0, 1, 'L')
        pdf.cell(90, 5, sanitize_text(f"A/C: {BANK_DETAILS['ac']}"), 0, 1, 'L')
        pdf.cell(90, 5, sanitize_text(f"IFSC: {BANK_DETAILS['ifsc']}"), 0, 1, 'L')
        pdf.cell(90, 5, sanitize_text(f"NAME: {BANK_DETAILS['name']}"), 0, 1, 'L')
        
        pdf.set_xy(130, y_foot + 15)
        pdf.set_font('Times', 'B', 10)
        pdf.cell(60, 5, sanitize_text(SIGNATORY_NAME), 0, 1, 'C')
        pdf.set_xy(130, y_foot + 20)
        pdf.cell(60, 5, "AUTHORIZED SIGNATORY", 0, 0, 'C')
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e: return None

# --- DOCX GENERATOR ---
def generate_docx_bytes(data, gst_rate_key, hide_gst, schedule_list, doc_no):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(10)
    ht = doc.add_table(rows=1, cols=2); ht.autofit = False; ht.columns[0].width = Inches(2.5); ht.columns[1].width = Inches(4.0)
    if os.path.exists(LOGO_FULL_PATH): 
        try: ht.cell(0,0).paragraphs[0].add_run().add_picture(LOGO_FULL_PATH, width=Inches(2.0))
        except: pass
    p = ht.cell(0,1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(COMPANY_NAME + "\n"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0, 0, 128)
    p.add_run(f"{COMPANY_ADDRESS}\nPh: {COMPANY_PHONE}")
    if not hide_gst: p.add_run(f"\nGST: {COMPANY_GSTIN}")
    doc.add_paragraph("_"*70)
    display_type = "BILL" if data['meta']['type'] == "FINAL BILL" else data['meta']['type']
    doc.add_paragraph(display_type).alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_table(rows=1, cols=2); t.autofit = True
    
    lbl = "Invoice No" if data['meta']['type'] == "FINAL BILL" else "Quotation No"
    t.cell(0,0).paragraphs[0].add_run(f"DETAILS:\n{lbl}: {doc_no}\nDate: {data['meta']['date']}")
    
    c_cell = t.cell(0,1)
    c_cell.paragraphs[0].add_run("TO CLIENT:\n").bold = True
    c_cell.paragraphs[0].add_run(f"{data['client']['name']}\n").bold = True
    c_cell.paragraphs[0].add_run(f"{data['client']['phone']}\n{data['client']['address']}")
    
    doc.add_paragraph("\n")
    # CHANGED: 4 Columns in Word Table
    tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Table Grid'
    hdrs = ["Description", "Qty", "Rate (Rs.)", "Amount (Rs.)"]
    for i,h in enumerate(hdrs): tbl.rows[0].cells[i].text = h
    for item in data['items']:
        rc = tbl.add_row().cells
        # Skipped Category Column
        rc[0].text=item['desc']
        rc[1].text=f"{item['qty']} {item['unit']}"; rc[2].text=f"{item['rate']:.2f}"; rc[3].text=f"{item['qty']*item['rate']:.2f}"
    sub, gst, grand = calculate_totals(data['items'], gst_rate_key)
    if hide_gst: gst=0; grand=sub
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"\nSubtotal: {sub:,.2f}\n"); 
    if not hide_gst: p.add_run(f"GST ({gst_rate_key}): {gst:,.2f}\n")
    p.add_run(f"Grand Total: Rs. {grand:,.2f}").bold = True
    p.add_run(f"\n{number_to_words_safe(grand)}").italic = True
    if schedule_list:
        doc.add_paragraph("\nPAYMENT SCHEDULE:").runs[0].bold = True
        stbl = doc.add_table(rows=1, cols=3); stbl.style = 'Table Grid'; shdrs = ["Stage", "Amount", "Date"]
        for i, h in enumerate(shdrs): stbl.rows[0].cells[i].text = h
        for row in schedule_list:
            rc = stbl.add_row().cells; rc[0].text=str(row.get("Stage","")); rc[1].text=str(row.get("Amount","")); rc[2].text=str(row.get("Date",""))
    doc.add_paragraph("\nTERMS:\n"+data['meta']['terms']); doc.add_paragraph("_"*70)
    ft = doc.add_table(rows=1, cols=2); ft.autofit = True
    ft.cell(0,0).paragraphs[0].add_run(f"BANK DETAILS\nBANK: {BANK_DETAILS['bank']}\nA/C: {BANK_DETAILS['ac']}\nIFSC: {BANK_DETAILS['ifsc']}")
    sig_cell = ft.cell(0,1).paragraphs[0]
    sig_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_cell.add_run(f"\n\n{SIGNATORY_NAME}\nAUTHORIZED SIGNATORY").bold = True
    f = io.BytesIO(); doc.save(f); f.seek(0); return f

# --- UI ---
st.markdown(f"<h1 style='color:#000080;'>🏗️ {COMPANY_NAME} Billing</h1>", unsafe_allow_html=True)

tab_b, tab_h, tab_t = st.tabs(["📝 Builder", "📂 History & Payments", "💰 Ledger"])

with tab_b:
    c1, c2, c3, c4 = st.columns(4)
    dtype = c1.selectbox("Type", ["QUOTATION", "FINAL BILL"], index=st.session_state.builder_dtype_idx)
    ddate = c2.date_input("Date", key="builder_date_picker")
    grate = c3.selectbox("GST", list(GST_RATES.keys()), index=3)
    hgst = c4.checkbox("Hide GST")
    
    st.markdown("---")
    cf, cp = st.columns([1, 1])
    
    with cf:
        c_name = st.text_input("Client Name", value=st.session_state.builder_c_name)
        c_mob = st.text_input("Client Mobile", value=st.session_state.builder_c_mob)
        c_addr = st.text_area("Client Address", value=st.session_state.builder_c_addr, height=60)
        
        st.session_state.builder_c_name = c_name
        st.session_state.builder_c_mob = c_mob
        st.session_state.builder_c_addr = c_addr
        
        st.subheader("Items")
        with st.container(border=True):
            # CHANGED: Description uses Work Catalog List
            descs = st.multiselect("Description", WORK_CATALOG)
            cust = st.text_input("Custom Desc.")
            c_q, c_r, c_u = st.columns(3)
            qty = c_q.number_input("Qty", 1.0)
            rate = c_r.number_input("Rate", 0.0, step=100.0)
            unit = c_u.selectbox("Unit", ["Sq.Ft", "Sq.Mt", "L/S", "Nos", "Job", "Sq.In", "Kg/Mt", "Secs"])
            
            if st.button("➕ Add"):
                d_list = descs[:]
                if cust: d_list.append(cust)
                if d_list:
                    # CHANGED: Category is empty string
                    st.session_state.invoice_data['items'].append({
                        "category": "", "desc": ", ".join(d_list), "unit": unit, "qty": qty, "rate": rate
                    })
                    st.rerun()
        
        if st.session_state.invoice_data['items']:
            st.dataframe(pd.DataFrame(st.session_state.invoice_data['items']), use_container_width=True)
            if st.button("Clear Items"): st.session_state.invoice_data['items'] = []; st.rerun()
            
        with st.expander("Payment Schedule (Optional)"):
            if 'Date' not in st.session_state.schedule_df.columns:
                st.session_state.schedule_df['Date'] = pd.Series(dtype='object')
                
            edited_sched = st.data_editor(
                st.session_state.schedule_df, 
                num_rows="dynamic", 
                use_container_width=True,
                column_config={
                    "Date": st.column_config.DateColumn("Date", format="YYYY-MM-DD")
                }
            )
            edited_sched['Date'] = edited_sched['Date'].apply(lambda x: x.strftime('%Y-%m-%d') if hasattr(x, 'strftime') else str(x) if pd.notnull(x) else "")
            
            st.session_state.schedule_df = edited_sched
            st.session_state.invoice_data['schedule'] = edited_sched.to_dict('records')

        term_txt = st.text_area("Terms", st.session_state.invoice_data['meta']['terms'], height=100)

    with cp:
        st.subheader("Live Preview")
        
        preview_id = generate_next_id(dtype, ddate)
        
        items = st.session_state.invoice_data['items']
        sub, gst, grand = calculate_totals(items, grate)
        if hgst: gst=0; grand=sub; gst_html=""
        else: gst_html=f"<tr><td colspan='3' align='right'>GST ({grate}):</td><td align='right'>Rs. {gst:,.2f}</td></tr>"
        
        logo_html = ""
        if os.path.exists(LOGO_FULL_PATH):
            with open(LOGO_FULL_PATH, "rb") as f: b64 = base64.b64encode(f.read()).decode()
            logo_html = f"<img src='data:image/png;base64,{b64}' width='120' style='vertical-align:top;'>"
            
        rows_str = ""
        for i in items:
            # CHANGED: HTML Table has 4 columns (Category Removed)
            rows_str += f"<tr><td>{i['desc']}</td><td align='right'>{i['qty']} {i['unit']}</td><td align='right'>{i['rate']}</td><td align='right'>{i['qty']*i['rate']:.2f}</td></tr>"
        
        schedule_html = ""
        sched_data = [r for r in st.session_state.invoice_data.get('schedule',[]) if r.get("Stage") or r.get("Amount")]
        if sched_data:
            sch_rows = "".join([f"<tr><td>{r.get('Stage','')}</td><td>{r.get('Amount','')}</td><td>{r.get('Date','')}</td></tr>" for r in sched_data])
            schedule_html = f"""<div style="margin-top:15px; border:1px solid #ccc;"><strong>PAYMENT SCHEDULE:</strong><table style="width:100%; border-collapse:collapse; font-size:12px;"><tr style="background:#eee;"><th>Stage</th><th>Amount</th><th>Date</th></tr>{sch_rows}</table></div>"""

        display_type_html = "BILL" if dtype == "FINAL BILL" else dtype
        lbl_preview = "Invoice No" if dtype == "FINAL BILL" else "Quotation No"

        html = f"""<div style="border:1px solid #ddd; padding:20px; font-family:'Times New Roman'; color:black; background:white;">
<table style="width:100%; border:none;"><tr><td style="width:50%; vertical-align:top;">{logo_html}</td><td style="width:50%; text-align:right; vertical-align:top;"><h2 style="color:#000080; margin:0;">{COMPANY_NAME}</h2><div style="font-size:12px; color:black;">{COMPANY_ADDRESS}<br>Ph: {COMPANY_PHONE}</div></td></tr></table>
<hr style="border: 1px solid #333; margin: 10px 0;"><h3 style="text-align:center;">{display_type_html}</h3>
<table style="width:100%; border-collapse:collapse; margin-bottom:20px;"><tr><td style="width:48%; border:1px solid #ccc; padding:10px; vertical-align:top;"><strong>DETAILS:</strong><br>Type: {display_type_html}<br>Date: {ddate}<br>{lbl_preview}: {preview_id}</td><td style="width:4%; border:none;"></td><td style="width:48%; border:1px solid #ccc; padding:10px; vertical-align:top;"><strong>TO CLIENT:</strong><br><strong style="font-size:16px;">{c_name}</strong><br>{c_mob}<br>{c_addr}</td></tr></table>
<table style="width:100%; border-collapse:collapse; border:1px solid #ccc; font-size:13px;" border="1"><tr style="background:#eee;"><th>Desc</th><th>Qty</th><th>Rate</th><th>Amt</th></tr>{rows_str}<tr><td colspan='3' align='right'>Subtotal:</td><td align='right'>Rs. {sub:,.2f}</td></tr>{gst_html}<tr><td colspan='3' align='right'><b>Total:</b></td><td align='right'><b>Rs. {grand:,.2f}</b></td></tr></table>
<p style="text-align:right; font-style:italic;">{number_to_words_safe(grand)}</p>
{schedule_html}
<div style="border:1px dashed #ccc; padding:10px; margin-top:10px;"><strong>TERMS:</strong><pre style="white-space:pre-wrap; font-family:inherit; margin:0;">{term_txt}</pre></div>
<div style="margin-top:20px; border-top:2px solid black; padding-top:10px; display:flex; justify-content:space-between;"><div><strong>BANK DETAILS</strong><br>BANK: {BANK_DETAILS['bank']}<br>A/C: {BANK_DETAILS['ac']}<br>IFSC: {BANK_DETAILS['ifsc']}</div>
<div style="text-align:center; min-width:200px;">
    <br>
    <div>{SIGNATORY_NAME}</div>
    <strong>AUTHORIZED SIGNATORY</strong>
</div>
</div></div>"""
        
        st.markdown(html, unsafe_allow_html=True)
        
        if st.button("💾 Finalize", type="primary"):
            if c_name:
                new_id = generate_next_id(dtype, ddate)
                
                rec = {
                    "id": f"INV-{int(datetime.now().timestamp())}",
                    "invoice_no": new_id,
                    "quotation_no": new_id if dtype == "QUOTATION" else None,
                    "date": str(ddate), 
                    "type": dtype, 
                    "client_name": c_name, 
                    "client_phone": c_mob,
                    "client_address": c_addr,
                    "amount": grand, 
                    "tax": gst, 
                    "items": items,
                    "gst_rate": grate,
                    "hide_gst": hgst,
                    "status": "Pending",
                    "schedule": st.session_state.invoice_data['schedule'],
                    "terms": term_txt
                }
                t = 'invoices' if dtype == "FINAL BILL" else 'quotations'
                st.session_state.db[t].append(rec); save_db(); st.toast(f"Saved: {new_id}")
            else: st.error("Name Required")
            
        fdata = {"meta": {"type": dtype, "date": str(ddate), "terms": term_txt}, "client": {"name": c_name, "phone": c_mob, "address": c_addr}, "items": items}
        
        pdf_bytes = generate_pdf_bytes(fdata, grate, hgst, sched_data, preview_id)
        f_suffix = "Bill" if dtype == "FINAL BILL" else "Quotation"
        
        if pdf_bytes:
            st.download_button("📄 Download PDF", pdf_bytes, f"{c_name} {f_suffix}.pdf", "application/pdf", type="primary")
        
        st.download_button("📝 Download Word", generate_docx_bytes(fdata, grate, hgst, sched_data, preview_id), f"{c_name} {f_suffix}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with tab_h:
    m = st.radio("View Mode", ["Quotations", "Bills & Payments"])
    
    if m == "Quotations":
        if st.session_state.db['quotations']:
            df_q = pd.DataFrame(st.session_state.db['quotations'])
            
            # CHANGED: Summary uses Desc instead of Category
            df_q['Items'] = df_q['items'].apply(lambda x: "; ".join([i['desc'] for i in x]))
            df_q.insert(0, 'S.No.', range(1, len(df_q) + 1))
            
            st.dataframe(df_q[['S.No.', 'quotation_no', 'date', 'client_name', 'amount', 'Items']], use_container_width=True, hide_index=True)
            
            sel_q_idx = st.selectbox("Select Quotation", range(len(df_q)), format_func=lambda x: f"{df_q.iloc[x]['client_name']} ({df_q.iloc[x]['quotation_no']})")
            selected_quote = df_q.iloc[sel_q_idx]
            
            c1, c2, c3 = st.columns(3)
            
            if c1.button("✅ Confirm as Bill"):
                st.session_state.invoice_data['items'] = selected_quote['items']
                if 'schedule' in selected_quote:
                    st.session_state.invoice_data['schedule'] = selected_quote['schedule']
                    st.session_state.schedule_df = pd.DataFrame(selected_quote['schedule'])
                
                st.session_state.builder_c_name = selected_quote['client_name']
                st.session_state.builder_c_mob = selected_quote.get('client_phone', '')
                st.session_state.builder_c_addr = selected_quote.get('client_address', '')
                st.session_state.builder_dtype_idx = 1
                
                st.success("Data loaded into Builder tab. Click 'Finalize' there to generate Bill No.")
            
            if c2.button("✏️ Edit Quote"):
                st.session_state.invoice_data['items'] = selected_quote['items']
                st.session_state.builder_c_name = selected_quote['client_name']
                st.session_state.builder_c_mob = selected_quote.get('client_phone', '')
                st.session_state.builder_c_addr = selected_quote.get('client_address', '')
                del st.session_state.db['quotations'][sel_q_idx]
                save_db()
                st.success("Loaded for Editing (Old deleted).")
                st.rerun()

            if c3.button("❌ Delete"):
                del st.session_state.db['quotations'][sel_q_idx]
                save_db()
                st.success("Deleted.")
                st.rerun()
                
            st.divider()
            st.write("📄 Download Copy:")
            fdata_h = {
                "meta": {"type": selected_quote['type'], "date": selected_quote['date'], "terms": selected_quote.get('terms', st.session_state.invoice_data['meta']['terms'])},
                "client": {"name": selected_quote['client_name'], "phone": selected_quote.get('client_phone',''), "address": selected_quote.get('client_address','')},
                "items": selected_quote['items']
            }
            
            doc_id = selected_quote.get('quotation_no', 'N/A')
            pdf_bytes_h = generate_pdf_bytes(fdata_h, selected_quote.get('gst_rate', '18%'), selected_quote.get('hide_gst', False), selected_quote.get('schedule', []), doc_id)
            
            if pdf_bytes_h:
                st.download_button("Download PDF", pdf_bytes_h, f"{selected_quote['client_name']} Quotation.pdf", "application/pdf")
            else:
                st.error("PDF Generation Failed")

        else: st.info("No active quotations.")

    else: 
        if st.session_state.db['invoices']:
            df_i = pd.DataFrame(st.session_state.db['invoices'])
            
            def get_paid(iid): return sum([p['amount'] for p in st.session_state.db['payments'] if p.get('invoice_id') == iid])
            
            df_i['Paid'] = df_i['id'].apply(get_paid)
            df_i['Pending'] = df_i['amount'] - df_i['Paid']
            
            for idx, row in df_i.iterrows():
                if row['Pending'] <= 0 and row['status'] != "Completed":
                    st.session_state.db['invoices'][idx]['status'] = "Completed"
                    save_db()
            
            df_i.insert(0, 'S.No.', range(1, len(df_i) + 1))
            
            if 'invoice_no' not in df_i.columns: df_i['invoice_no'] = df_i['id']
            
            st.dataframe(df_i[['S.No.', 'invoice_no', 'date', 'client_name', 'amount', 'Paid', 'Pending', 'status']], use_container_width=True, hide_index=True)
            
            st.divider()
            
            # --- DOWNLOAD BILL FROM HISTORY ---
            sel_b_idx = st.selectbox("Select Bill", range(len(df_i)), format_func=lambda x: f"{df_i.iloc[x]['client_name']} ({df_i.iloc[x]['invoice_no']})")
            db_record = st.session_state.db['invoices'][sel_b_idx]
            selected_bill_df = df_i.iloc[sel_b_idx]
            
            c1, c2 = st.columns(2)
            
            fdata_b = {
                "meta": {"type": db_record['type'], "date": db_record['date'], "terms": db_record.get('terms', "")},
                "client": {"name": db_record['client_name'], "phone": db_record.get('client_phone',''), "address": db_record.get('client_address','')},
                "items": db_record['items']
            }
            
            doc_id = db_record.get('invoice_no', 'N/A')
            pdf_bytes_b = generate_pdf_bytes(fdata_b, db_record.get('gst_rate','18%'), db_record.get('hide_gst', False), db_record.get('schedule',[]), doc_id)
            
            if pdf_bytes_b:
                c1.download_button("📄 Download Bill PDF", pdf_bytes_b, f"{db_record['client_name']} Bill.pdf", "application/pdf", type="primary")
            else:
                c1.error("PDF Generation Failed")
            
            if db_record['status'] == "Pending":
                if c2.button("✅ Mark as Complete (Force)"):
                    st.session_state.db['invoices'][sel_b_idx]['status'] = "Completed (Manual)"
                    save_db()
                    st.rerun()

            st.markdown("### Payment History & Receipts")
            
            bill_payments = [p for p in st.session_state.db['payments'] if p.get('invoice_id') == db_record['id']]
            if bill_payments:
                for idx, p in enumerate(bill_payments):
                    pc1, pc2, pc3, pc4 = st.columns([2,2,2,2])
                    pc1.write(f"**Date:** {p['date']}")
                    pc2.write(f"**Amt:** {p['amount']}")
                    pc3.write(f"**Mode:** {p['mode']}")
                    
                    rec_bytes = generate_receipt_bytes(p)
                    if rec_bytes:
                        pc4.download_button(f"📥 Receipt {idx+1}", rec_bytes, f"Receipt_{p['date']}.pdf", "application/pdf", key=f"d_rec_{idx}")
                    else:
                        pc4.error("Err")
            else:
                st.caption("No payments recorded yet.")

            st.divider()

            if db_record['status'] == "Pending" and selected_bill_df['Pending'] > 1.0:
                st.subheader("💰 Record New Payment")
                with st.container(border=True):
                    c1, c2, c3 = st.columns(3)
                    pay_amt = c1.number_input("Amount", max_value=float(selected_bill_df['Pending']), value=float(selected_bill_df['Pending']), key="pay_amt")
                    pay_date = c2.date_input("Date", key="pay_date_picker")
                    pay_mode = c3.selectbox("Mode", ["UPI", "Cash", "Cheque", "Transfer"], key="pay_mode_sel")
                    
                    if st.button("Save Payment", type="primary"):
                        p_rec = {
                            "id": f"PAY-{int(datetime.now().timestamp())}",
                            "invoice_id": db_record['id'],
                            "client_name": db_record['client_name'],
                            "invoice_date": db_record['date'],
                            "amount": pay_amt,
                            "date": str(pay_date),
                            "mode": pay_mode
                        }
                        st.session_state.db['payments'].append(p_rec)
                        save_db()
                        st.session_state.last_pay = p_rec
                        st.success("Payment Recorded!")
                        st.rerun()

            else:
                st.success("✅ This bill is Fully Paid / Completed.")

with tab_t:
    st.header("Financial Ledger")
    
    # Date Filter
    c_d1, c_d2 = st.columns(2)
    d_from = c_d1.date_input("From Date", date(date.today().year, 1, 1))
    d_to = c_d2.date_input("To Date", date.today())
    
    invs = st.session_state.db.get('invoices', [])
    pays = st.session_state.db.get('payments', [])
    
    def parse_date(d_str):
        try: return datetime.strptime(d_str, '%Y-%m-%d').date()
        except: return date.min
    
    invs_filtered = [i for i in invs if d_from <= parse_date(i['date']) <= d_to]
    pays_filtered = [p for p in pays if d_from <= parse_date(p['date']) <= d_to]
    
    t_billed = sum(i['amount'] for i in invs_filtered)
    t_rev = sum(p['amount'] for p in pays_filtered)
    t_gst = sum(i['tax'] for i in invs_filtered)
    
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Total Work Billed", f"Rs. {t_billed:,.2f}")
    m2.metric("Actual Revenue", f"Rs. {t_rev:,.2f}")
    m3.metric("Pending Dues (Total)", f"Rs. {sum(i['amount'] for i in invs) - sum(p['amount'] for p in pays):,.2f}")
    m4.metric("GST Liability", f"Rs. {t_gst:,.2f}")
    
    st.divider()
    st.subheader("Client Reports")
    clients = sorted(list(set([i['client_name'] for i in invs])))
    sel_c = st.selectbox("Select Client", ["All"] + clients)
    
    def clean_df(data_list, is_invoice=True):
        if not data_list: return pd.DataFrame()
        df = pd.DataFrame(data_list)
        df.insert(0, 'S.No.', range(1, len(df) + 1))
        # CHANGED: Items use desc, not category
        if 'items' in df.columns:
            df['Items'] = df['items'].apply(lambda x: "; ".join([i['desc'] for i in x]) if isinstance(x, list) else "")
        if is_invoice and 'invoice_no' not in df.columns: df['invoice_no'] = df.get('id', '')
        
        if is_invoice:
            cols = ['S.No.', 'invoice_no', 'date', 'client_name', 'amount', 'Items']
        else:
            cols = ['S.No.', 'date', 'amount', 'mode', 'client_name']
            
        return df[[c for c in cols if c in df.columns]]

    if sel_c != "All":
        c_inv = [i for i in invs_filtered if i['client_name'] == sel_c]
        c_pay = [p for p in pays_filtered if p['client_name'] == sel_c]
        st.write(f"**Total Billed:** {sum(x['amount'] for x in c_inv):,.2f} | **Total Paid:** {sum(x['amount'] for x in c_pay):,.2f}")
        
        st.write("Bill History (Filtered):")
        st.dataframe(clean_df(c_inv, True), use_container_width=True, hide_index=True)
        
        st.write("Payment History (Filtered):")
        st.dataframe(clean_df(c_pay, False), use_container_width=True, hide_index=True)
    else:
        st.write("All Bills (Filtered):")
        st.dataframe(clean_df(invs_filtered, True), use_container_width=True, hide_index=True)
        
        st.write("All Payments (Filtered):")
        st.dataframe(clean_df(pays_filtered, False), use_container_width=True, hide_index=True)
        
    st.divider()
    c1, c2 = st.columns(2)
    if invs_filtered: c1.download_button("📥 Export Bills CSV", pd.DataFrame(invs_filtered).to_csv(), "bills_filtered.csv")
    if pays_filtered: c2.download_button("📥 Export Revenue CSV", pd.DataFrame(pays_filtered).to_csv(), "revenue_filtered.csv")
